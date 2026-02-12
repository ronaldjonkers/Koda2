"""Document and code generation service."""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any, Optional

from jinja2 import Environment, FileSystemLoader, BaseLoader

from executiveai.logging_config import get_logger

logger = get_logger(__name__)


class DocumentService:
    """Generate DOCX, XLSX, PDF documents and scaffold code projects."""

    def __init__(self, template_dir: str = "templates") -> None:
        self._template_dir = Path(template_dir)
        self._template_dir.mkdir(parents=True, exist_ok=True)
        self._jinja_env = Environment(
            loader=FileSystemLoader(str(self._template_dir)),
            autoescape=False,
        )

    # ── DOCX ─────────────────────────────────────────────────────────

    def generate_docx(
        self,
        title: str,
        content: list[dict[str, Any]],
        output_path: str,
        template_path: Optional[str] = None,
    ) -> str:
        """Generate a Word document.

        Args:
            title: Document title.
            content: List of dicts with keys 'type' (heading/paragraph/table) and 'data'.
            output_path: Where to save the document.
            template_path: Optional .docx template to use as base.
        """
        from docx import Document
        from docx.shared import Inches, Pt

        doc = Document(template_path) if template_path else Document()
        doc.add_heading(title, level=0)

        for block in content:
            block_type = block.get("type", "paragraph")
            data = block.get("data", "")

            if block_type == "heading":
                level = block.get("level", 1)
                doc.add_heading(str(data), level=level)
            elif block_type == "paragraph":
                doc.add_paragraph(str(data))
            elif block_type == "table":
                rows = data if isinstance(data, list) else []
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = "Table Grid"
                    for i, row in enumerate(rows):
                        for j, cell in enumerate(row):
                            table.cell(i, j).text = str(cell)
            elif block_type == "image":
                doc.add_picture(str(data), width=Inches(5))
            elif block_type == "page_break":
                doc.add_page_break()

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        logger.info("docx_generated", path=output_path)
        return output_path

    # ── XLSX ─────────────────────────────────────────────────────────

    def generate_xlsx(
        self,
        title: str,
        sheets: dict[str, list[list[Any]]],
        output_path: str,
    ) -> str:
        """Generate an Excel workbook with multiple sheets.

        Args:
            title: Workbook title (used as first sheet name if only one sheet).
            sheets: Dict of sheet_name -> rows (list of lists).
            output_path: Where to save the workbook.
        """
        from openpyxl import Workbook
        from openpyxl.styles import Font

        wb = Workbook()
        first = True
        for sheet_name, rows in sheets.items():
            ws = wb.active if first else wb.create_sheet(title=sheet_name)
            if first:
                ws.title = sheet_name
                first = False
            for r_idx, row in enumerate(rows, 1):
                for c_idx, value in enumerate(row, 1):
                    cell = ws.cell(row=r_idx, column=c_idx, value=value)
                    if r_idx == 1:
                        cell.font = Font(bold=True)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        wb.save(output_path)
        logger.info("xlsx_generated", path=output_path)
        return output_path

    # ── PDF ──────────────────────────────────────────────────────────

    def generate_pdf(
        self,
        title: str,
        content: list[dict[str, Any]],
        output_path: str,
    ) -> str:
        """Generate a PDF document using ReportLab.

        Args:
            title: Document title.
            content: List of dicts with 'type' and 'data'.
            output_path: Where to save the PDF.
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        from reportlab.lib import colors

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 0.3 * inch))

        for block in content:
            block_type = block.get("type", "paragraph")
            data = block.get("data", "")

            if block_type == "heading":
                level = block.get("level", 1)
                style = styles[f"Heading{min(level, 6)}"] if level <= 6 else styles["Heading6"]
                story.append(Paragraph(str(data), style))
            elif block_type == "paragraph":
                story.append(Paragraph(str(data), styles["Normal"]))
                story.append(Spacer(1, 0.1 * inch))
            elif block_type == "table":
                rows = data if isinstance(data, list) else []
                if rows:
                    t = Table(rows)
                    t.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 0.2 * inch))

        doc.build(story)
        logger.info("pdf_generated", path=output_path)
        return output_path

    # ── Template Rendering ───────────────────────────────────────────

    def render_template(self, template_name: str, variables: dict[str, Any]) -> str:
        """Render a Jinja2 template from the template directory."""
        tmpl = self._jinja_env.get_template(template_name)
        return tmpl.render(**variables)

    def render_string_template(self, template_str: str, variables: dict[str, Any]) -> str:
        """Render a Jinja2 template from a string."""
        env = Environment(loader=BaseLoader(), autoescape=False)
        tmpl = env.from_string(template_str)
        return tmpl.render(**variables)

    # ── Code Generation / Scaffolding ────────────────────────────────

    def scaffold_project(
        self,
        project_name: str,
        project_type: str,
        output_dir: str,
        options: Optional[dict[str, Any]] = None,
    ) -> str:
        """Scaffold a new code project.

        Args:
            project_name: Name of the project.
            project_type: Type (python, node, fastapi, react).
            output_dir: Directory to create the project in.
            options: Additional scaffolding options.
        """
        opts = options or {}
        base = Path(output_dir) / project_name
        base.mkdir(parents=True, exist_ok=True)

        scaffolds: dict[str, dict[str, str]] = {
            "python": {
                "pyproject.toml": f'[project]\nname = "{project_name}"\nversion = "0.1.0"\n',
                f"{project_name}/__init__.py": f'"""Package {project_name}."""\n',
                f"{project_name}/main.py": 'def main():\n    print("Hello")\n\nif __name__ == "__main__":\n    main()\n',
                "tests/__init__.py": "",
                "tests/test_main.py": f"from {project_name}.main import main\n\ndef test_main():\n    main()\n",
                "README.md": f"# {project_name}\n",
                ".gitignore": "__pycache__/\n*.pyc\n.venv/\n",
            },
            "fastapi": {
                "pyproject.toml": f'[project]\nname = "{project_name}"\nversion = "0.1.0"\ndependencies = ["fastapi", "uvicorn"]\n',
                "app/__init__.py": "",
                "app/main.py": 'from fastapi import FastAPI\n\napp = FastAPI()\n\n@app.get("/")\ndef root():\n    return {"message": "Hello"}\n',
                "tests/__init__.py": "",
                "README.md": f"# {project_name}\n",
                "Dockerfile": 'FROM python:3.12-slim\nWORKDIR /app\nCOPY . .\nRUN pip install .\nCMD ["uvicorn", "app.main:app"]\n',
            },
        }

        template = scaffolds.get(project_type, scaffolds["python"])
        for filepath, content in template.items():
            full_path = base / filepath
            full_path.parent.mkdir(parents=True, exist_ok=True)
            full_path.write_text(content)

        logger.info("project_scaffolded", name=project_name, type=project_type, path=str(base))
        return str(base)
